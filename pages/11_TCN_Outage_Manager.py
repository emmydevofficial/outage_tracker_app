"""
### FILE: pages/11_TCN_Outage_Manager.py
TCN Outage Attributes Manager — merged into the main multipage app.

Ports the standalone TCN-Outage-Manager-main Streamlit app in as a single
page. Data here is file-based (Excel/CSV/Word attribute files dropped into
`tcn_uploads/`), separate from the PostgreSQL-backed pages elsewhere in this
app. Authentication reuses the shared DB-backed login (utils.auth.login)
instead of the original app's local users.json file.
"""

import streamlit as st
from utils.auth import login, filter_to_user_region, scoped_regions, is_super_admin, current_region
from utils.activity_log import log_activity
from utils.file_storage import save_uploaded_file

login()

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from pathlib import Path
from datetime import datetime, timedelta, date
import re
import io
import base64
from equipment_data import (
    SUBSTATION_EQUIPMENT, EQUIPMENT_TYPES,
    get_substations_for, get_equipment_types_for, get_equipment_list,
)

st.set_page_config(
    page_title="TCN Outage Attributes Manager",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Logo + Background styling ──────────────────────────────────────────────
ASSETS_DIR = Path(__file__).parent.parent / "assets"
_logo_path = ASSETS_DIR / "tcn_logo.png"
_logo_b64 = base64.b64encode(_logo_path.read_bytes()).decode() if _logo_path.exists() else ""
_bg_path = ASSETS_DIR / "tcn_background.png"
if _bg_path.exists():
    _bg_b64 = base64.b64encode(_bg_path.read_bytes()).decode()
    st.markdown(f"""<style>
    /* ── Impeccable Design Tokens ──────────────────────────────────── */
    /* Product register: Restrained color strategy. OKLCH tinted neutrals. */
    /* No pure black/white. Neutrals tinted toward brand blue hue (250). */
    :root {{
        /* Easing: ease-out-quart for natural deceleration (Emil + Impeccable) */
        --ease-out: cubic-bezier(0.25, 1, 0.5, 1);
        --ease-out-quart: cubic-bezier(0.165, 0.84, 0.44, 1);
        --ease-in-out: cubic-bezier(0.77, 0, 0.175, 1);

        /* Tinted neutrals toward brand blue (hue 250) */
        --neutral-950: oklch(12% 0.01 250);
        --neutral-900: oklch(18% 0.008 250);
        --neutral-700: oklch(35% 0.008 250);
        --neutral-500: oklch(55% 0.006 250);
        --neutral-300: oklch(78% 0.005 250);
        --neutral-200: oklch(88% 0.004 250);
        --neutral-100: oklch(94% 0.003 250);
        --neutral-50: oklch(97.5% 0.002 250);

        /* Brand: TCN institutional blue + signal red */
        --brand-primary: oklch(32% 0.12 250);
        --brand-primary-light: oklch(42% 0.10 250);
        --brand-accent: oklch(48% 0.18 25);
        --brand-accent-hover: oklch(42% 0.20 25);

        /* Semantic */
        --surface-primary: oklch(97.5% 0.002 250);
        --surface-elevated: oklch(99% 0.002 250);
        --surface-sidebar: oklch(28% 0.10 250);
        --border-subtle: oklch(88% 0.005 250);
        --border-sidebar-input: oklch(50% 0.04 250);
        --text-primary: oklch(18% 0.008 250);
        --text-secondary: oklch(45% 0.006 250);
        --text-on-dark: oklch(95% 0.003 250);
    }}

    /* ── Base App ──────────────────────────────────────────────────── */
    .stApp {{
        background-image: url("data:image/png;base64,{_bg_b64}");
        background-size: cover;
        background-position: center;
        background-attachment: fixed;
        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", system-ui, sans-serif;
    }}
    .stApp > header {{
        background: transparent;
    }}

    /* ── Sidebar ───────────────────────────────────────────────────── */
    /* No side-stripe border (Impeccable absolute ban). Full surface color. */
    [data-testid="stSidebar"] {{
        background: var(--surface-sidebar);
    }}
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] .stMarkdown,
    [data-testid="stSidebar"] .stMarkdown p,
    [data-testid="stSidebar"] .stMarkdown strong,
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3,
    [data-testid="stSidebar"] [data-testid="stWidgetLabel"] {{
        color: var(--text-on-dark) !important;
    }}

    /* Sidebar date inputs */
    [data-testid="stSidebar"] [data-testid="stDateInput"] input {{
        background: transparent !important;
        color: var(--text-on-dark) !important;
    }}
    [data-testid="stSidebar"] [data-testid="stDateInput"] [data-baseweb="input"] {{
        background: oklch(35% 0.06 250) !important;
        border: 1px solid var(--border-sidebar-input) !important;
    }}
    [data-testid="stSidebar"] [data-testid="stDateInput"] [data-baseweb="input"] * {{
        color: var(--text-on-dark) !important;
    }}

    /* Sidebar search input */
    [data-testid="stSidebar"] [data-testid="stTextInputRootElement"] {{
        background: oklch(35% 0.06 250) !important;
        border: 1px solid var(--border-sidebar-input) !important;
        border-radius: 6px;
    }}
    [data-testid="stSidebar"] [data-testid="stTextInputRootElement"] input {{
        background: transparent !important;
        color: var(--text-on-dark) !important;
    }}
    [data-testid="stSidebar"] [data-testid="stTextInputRootElement"] input::placeholder {{
        color: oklch(70% 0.01 250) !important;
    }}
    [data-testid="stSidebar"] [data-testid="stTextInputRootElement"] [data-baseweb="base-input"],
    [data-testid="stSidebar"] [data-testid="stTextInputRootElement"] [data-baseweb="input"] {{
        background: transparent !important;
        border: none !important;
    }}

    /* Sidebar logout button */
    [data-testid="stSidebar"] .stButton button {{
        background: var(--brand-accent) !important;
        color: var(--text-on-dark) !important;
        border: none !important;
        font-weight: 500;
    }}
    [data-testid="stSidebar"] .stButton button:hover {{
        background: var(--brand-accent-hover) !important;
    }}

    /* Sidebar multiselect */
    [data-testid="stSidebar"] .stMultiSelect [data-baseweb="select"] {{
        background: oklch(95% 0.003 250) !important;
    }}
    [data-testid="stSidebar"] .stMultiSelect [data-baseweb="tag"] {{
        background: var(--brand-primary) !important;
    }}
    [data-testid="stSidebar"] .stMultiSelect [data-baseweb="tag"] span {{
        color: var(--text-on-dark) !important;
    }}
    [data-testid="stSidebar"] .stMultiSelect input {{
        color: var(--text-primary) !important;
    }}
    [data-testid="stSidebar"] [data-baseweb="popover"] * {{
        color: var(--text-primary) !important;
    }}

    /* ── Tabs ──────────────────────────────────────────────────────── */
    .stTabs [data-baseweb="tab-list"] {{
        background: oklch(96% 0.002 250);
        border: 1px solid oklch(90% 0.004 250);
        border-radius: 8px;
        padding: 3px;
    }}
    .stTabs [data-baseweb="tab"] {{
        color: var(--text-secondary);
        font-weight: 500;
        border-radius: 6px;
        transition: color 150ms var(--ease-out-quart),
                    background 150ms var(--ease-out-quart) !important;
    }}
    .stTabs [data-baseweb="tab"]:hover {{
        color: var(--brand-primary) !important;
        background: oklch(98% 0.002 250);
    }}
    .stTabs [aria-selected="true"] {{
        color: var(--brand-accent) !important;
        background: var(--surface-elevated) !important;
        border-bottom: 2px solid var(--brand-accent);
        font-weight: 600;
        box-shadow: 0 1px 3px oklch(32% 0.06 250 / 0.06);
    }}
    .stTabs [data-baseweb="tab"]:active {{
        transform: scale(0.98);
    }}

    /* ── Forms (Double-Bezel) ─────────────────────────────────────── */
    /* No glassmorphism (Impeccable absolute ban). Nested shell architecture. */
    [data-testid="stForm"] {{
        background: oklch(96% 0.002 250);
        border: 1px solid oklch(90% 0.004 250);
        border-radius: 12px;
        padding: 4px;
        transition: box-shadow 250ms var(--ease-out-quart) !important;
    }}
    [data-testid="stForm"] > div {{
        background: var(--surface-elevated);
        border-radius: 9px;
        padding: 1.5rem;
        box-shadow: inset 0 1px 1px oklch(100% 0 0 / 0.5);
    }}
    [data-testid="stForm"]:hover {{
        box-shadow: 0 4px 20px oklch(32% 0.06 250 / 0.05),
                    0 1px 3px oklch(32% 0.06 250 / 0.03) !important;
    }}

    /* ── Metric Cards (Double-Bezel Architecture) ───────────────── */
    /* Outer shell + inner core for machined-hardware depth feel. */
    .stMetric {{
        background: oklch(96% 0.002 250);
        border: 1px solid oklch(90% 0.004 250);
        border-radius: 10px;
        padding: 3px;
        transition: box-shadow 250ms var(--ease-out-quart) !important;
    }}
    .stMetric > div {{
        background: var(--surface-elevated);
        border-radius: 8px;
        padding: 0.75rem 1rem;
        box-shadow: inset 0 1px 1px oklch(100% 0 0 / 0.6);
    }}
    .stMetric:hover {{
        box-shadow: 0 4px 20px oklch(32% 0.06 250 / 0.06),
                    0 1px 3px oklch(32% 0.06 250 / 0.04) !important;
    }}
    .stMetric label {{
        color: var(--text-secondary) !important;
        font-weight: 500;
        font-size: 0.75rem;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }}
    .stMetric [data-testid="stMetricValue"] {{
        color: var(--text-primary) !important;
        font-weight: 600;
        font-family: "SF Mono", "Geist Mono", "JetBrains Mono", "Cascadia Code", ui-monospace, monospace !important;
        letter-spacing: -0.01em;
    }}
    /* Staggered entry animation for metric cards */
    @keyframes metricFadeUp {{
        from {{ opacity: 0; transform: translateY(10px); }}
        to {{ opacity: 1; transform: translateY(0); }}
    }}
    [data-testid="stHorizontalBlock"] > div:nth-child(1) .stMetric {{ animation: metricFadeUp 500ms var(--ease-out-quart) both; animation-delay: 0ms; }}
    [data-testid="stHorizontalBlock"] > div:nth-child(2) .stMetric {{ animation: metricFadeUp 500ms var(--ease-out-quart) both; animation-delay: 60ms; }}
    [data-testid="stHorizontalBlock"] > div:nth-child(3) .stMetric {{ animation: metricFadeUp 500ms var(--ease-out-quart) both; animation-delay: 120ms; }}
    [data-testid="stHorizontalBlock"] > div:nth-child(4) .stMetric {{ animation: metricFadeUp 500ms var(--ease-out-quart) both; animation-delay: 180ms; }}
    [data-testid="stHorizontalBlock"] > div:nth-child(5) .stMetric {{ animation: metricFadeUp 500ms var(--ease-out-quart) both; animation-delay: 240ms; }}

    /* ── Main Content Focus States ────────────────────────────────── */
    [data-testid="stTextInputRootElement"]:focus-within,
    [data-testid="stNumberInputContainer"]:focus-within,
    [data-baseweb="select"]:focus-within {{
        outline: 2px solid oklch(50% 0.10 250 / 0.35);
        outline-offset: 1px;
        border-radius: 6px;
        transition: outline 150ms var(--ease-out-quart) !important;
    }}

    /* ── Plotly Modebar ────────────────────────────────────────────── */
    /* Hide modebar by default; show on chart hover for clean presentation. */
    .modebar {{
        opacity: 0 !important;
        transition: opacity 200ms var(--ease-out-quart) !important;
    }}
    [data-testid="stPlotlyChart"]:hover .modebar {{
        opacity: 1 !important;
    }}

    /* ── Buttons: Press Feedback ───────────────────────────────────── */
    /* 150-250ms, ease-out-quart. scale(0.97) on :active. */
    .stButton button {{
        transition: transform 150ms var(--ease-out-quart),
                    box-shadow 200ms var(--ease-out-quart),
                    background 200ms var(--ease-out-quart) !important;
    }}
    .stButton button:active {{
        transform: scale(0.97) !important;
    }}
    .stButton button:hover {{
        box-shadow: 0 2px 12px oklch(32% 0.06 250 / 0.08) !important;
    }}
    [data-testid="stSidebar"] .stButton button:active {{
        transform: scale(0.97) !important;
    }}

    /* ── Download Buttons ──────────────────────────────────────────── */
    .stDownloadButton button {{
        transition: transform 150ms var(--ease-out-quart),
                    box-shadow 200ms var(--ease-out-quart) !important;
    }}
    .stDownloadButton button:active {{
        transform: scale(0.97) !important;
    }}

    /* ── Sidebar Focus States ──────────────────────────────────────── */
    [data-testid="stSidebar"] [data-testid="stTextInputRootElement"]:focus-within {{
        border-color: oklch(60% 0.08 250) !important;
        box-shadow: 0 0 0 2px oklch(50% 0.08 250 / 0.2) !important;
        transition: border-color 150ms var(--ease-out-quart),
                    box-shadow 150ms var(--ease-out-quart) !important;
    }}
    [data-testid="stSidebar"] [data-testid="stDateInput"] [data-baseweb="input"]:focus-within {{
        border-color: oklch(60% 0.08 250) !important;
        box-shadow: 0 0 0 2px oklch(50% 0.08 250 / 0.2) !important;
        transition: border-color 150ms var(--ease-out-quart),
                    box-shadow 150ms var(--ease-out-quart) !important;
    }}

    /* ── Multiselect Tags ──────────────────────────────────────────── */
    [data-testid="stSidebar"] .stMultiSelect [data-baseweb="tag"] {{
        transition: opacity 100ms var(--ease-out-quart) !important;
    }}

    /* ── Data Tables (Double-Bezel) ────────────────────────────────── */
    [data-testid="stDataFrame"] {{
        background: oklch(96% 0.002 250);
        border: 1px solid oklch(90% 0.004 250);
        border-radius: 10px;
        padding: 3px;
        overflow: hidden;
        transition: box-shadow 250ms var(--ease-out-quart) !important;
    }}
    [data-testid="stDataFrame"] > div {{
        border-radius: 8px;
        overflow: hidden;
        box-shadow: inset 0 1px 1px oklch(100% 0 0 / 0.4);
    }}
    [data-testid="stDataFrame"]:hover {{
        box-shadow: 0 4px 20px oklch(32% 0.06 250 / 0.05),
                    0 1px 3px oklch(32% 0.06 250 / 0.03) !important;
    }}

    /* ── Chart Containers (Double-Bezel) ─────────────────────────── */
    [data-testid="stPlotlyChart"] {{
        background: oklch(96% 0.002 250);
        border: 1px solid oklch(90% 0.004 250);
        border-radius: 12px;
        padding: 4px;
        transition: box-shadow 250ms var(--ease-out-quart) !important;
    }}
    [data-testid="stPlotlyChart"] > div {{
        background: var(--surface-elevated);
        border-radius: 9px;
        box-shadow: inset 0 1px 1px oklch(100% 0 0 / 0.5);
    }}
    [data-testid="stPlotlyChart"]:hover {{
        box-shadow: 0 4px 24px oklch(32% 0.06 250 / 0.05),
                    0 1px 4px oklch(32% 0.06 250 / 0.03) !important;
    }}

    /* ── Expanders ─────────────────────────────────────────────────── */
    [data-testid="stExpander"] {{
        border-radius: 10px;
        overflow: hidden;
        transition: box-shadow 250ms var(--ease-out-quart) !important;
    }}
    [data-testid="stExpander"]:hover {{
        box-shadow: 0 4px 16px oklch(32% 0.06 250 / 0.05) !important;
    }}

    /* ── File Uploader ─────────────────────────────────────────────── */
    [data-testid="stFileUploader"] {{
        border-radius: 8px;
    }}

    /* ── Skeleton Shimmer (Loading States) ─────────────────────────── */
    @keyframes shimmer {{
        0% {{ background-position: -200% 0; }}
        100% {{ background-position: 200% 0; }}
    }}
    .stSpinner > div {{
        background: linear-gradient(
            90deg,
            oklch(95% 0.002 250) 25%,
            oklch(90% 0.003 250) 50%,
            oklch(95% 0.002 250) 75%
        );
        background-size: 200% 100%;
        animation: shimmer 1.8s var(--ease-in-out) infinite;
        border-radius: 6px;
    }}

    /* ── Scrollbar (Refined) ───────────────────────────────────────── */
    ::-webkit-scrollbar {{
        width: 6px;
        height: 6px;
    }}
    ::-webkit-scrollbar-track {{
        background: transparent;
    }}
    ::-webkit-scrollbar-thumb {{
        background: oklch(78% 0.005 250);
        border-radius: 3px;
    }}
    ::-webkit-scrollbar-thumb:hover {{
        background: oklch(65% 0.006 250);
    }}
    [data-testid="stSidebar"] ::-webkit-scrollbar-thumb {{
        background: oklch(45% 0.04 250);
    }}

    /* ── Accessibility: Reduced Motion ─────────────────────────────── */
    @media (prefers-reduced-motion: reduce) {{
        *, *::before, *::after {{
            transition-duration: 0.01ms !important;
            animation-duration: 0.01ms !important;
        }}
        .stMetric:hover {{
            transform: none !important;
        }}
    }}
    </style>""", unsafe_allow_html=True)

DATA_DIR = Path(__file__).parent.parent
UPLOADS_DIR = DATA_DIR / "tcn_uploads"
UPLOADS_DIR.mkdir(exist_ok=True)
SOURCE_COLUMNS = [
    "Disco", "Region", "SubRegion_ACC", "Substation", "Feeder_33kV",
    "Date_Off", "Hour_Off", "Minute_Off", "Date_On", "Hour_On", "Minute_On",
    "Duration", "Class", "Last_Load_MW", "Event_Indication",
    "Officer_Interruption", "Officer_Restoration",
    "Party_Responsible", "Weather_Condition", "Remarks",
]
_no_party = [c for c in SOURCE_COLUMNS if c != "Party_Responsible"]
COLUMNS = _no_party[:14] + ["Frequency_Hz"] + _no_party[14:]
DISPLAY_COLUMNS = [
    "Disco", "Region", "SubRegion/ACC", "Substation", "33kV Feeder",
    "Date Off", "Hour Off", "Minute Off", "Date On", "Hour On", "Minute On",
    "Duration (H:mm)", "Class", "Last Load (MW)", "Frequency (Hz)", "Event/Indication",
    "Officer (Interruption)", "Officer (Restoration)",
    "Weather", "Remarks",
]

DISCOS = [
    "Abuja", "Benin", "Bilateral", "Eko", "Enugu",
    "Ibadan", "Ikeja", "Jos", "Kaduna", "Kano", "Port-Harcourt", "Yola",
]

REGION_DISCO_MAP = {
    "Abuja": ["Abuja", "Benin", "Jos"],
    "Bauchi": ["Jos", "Yola"],
    "Benin": ["Benin"],
    "Enugu": ["Enugu"],
    "Kaduna": ["Kaduna"],
    "Kano": ["Kano"],
    "Lagos": ["Eko", "Ikeja"],
    "Osogbo": ["Ibadan"],
    "Port-Harcourt": ["Port-Harcourt"],
    "Shiroro": ["Abuja", "Kaduna", "Kano"],
}

EVENT_INDICATIONS = [
    "Line Trip", "E/F", "O/C", "U/V", "O/V", "Buchholz",
    "Diff. Protection", "Over Fluxing", "Restricted E/F",
    "Surge Arrester", "Transformer Trip", "Bus Bar Protection",
    "Distance Protection", "Power Swings", "Load Shedding",
    "Planned Outage", "Emergency Outage", "System Collapse",
    "Feeder Trip", "Equipment Failure", "Vandalism",
    "Fire Outbreak", "Flooding", "Vegetation/Bush Burning",
    "Jumper Cut", "Broken Conductor", "Tower Collapse",
    "Insulator Failure", "Oil Leakage", "Expired CT/VT",
    "Relay Maloperation", "Manual Trip", "Frequency Relay",
    "Other",
]

OUTAGE_CLASSES = ["Forced", "Emergency", "Planned"]
WEATHER_CONDITIONS = ["Clear", "Rain", "Cloudy", "Storm", "Harmattan", "Fine", "Hazy", "Foggy"]


def parse_excel_file(fpath):
    rows = []
    try:
        xls = pd.ExcelFile(fpath, engine="openpyxl")
        for sheet in xls.sheet_names:
            df = pd.read_excel(fpath, sheet_name=sheet, header=None, engine="openpyxl")
            header_row = None
            for idx, row in df.iterrows():
                vals = [str(v).strip().lower() for v in row.values[:5]]
                if "disco" in vals and "region" in vals:
                    header_row = idx
                    break
            if header_row is None:
                continue
            data = df.iloc[header_row + 1:].copy()
            ncols = data.shape[1]
            if ncols > 20:
                data = data.iloc[:, :20]
            elif ncols < 20:
                for i in range(20 - ncols):
                    data[f"_pad_{i}"] = None
            data.columns = SOURCE_COLUMNS
            data = data.dropna(subset=["Disco", "Region"], how="all")
            data = data[~data["Disco"].astype(str).str.contains("NO TCN ATTRIBUTION", case=False, na=False)]
            data = data[data["Disco"].notna() & (data["Disco"].astype(str).str.strip() != "")]
            data = data[data["Disco"].astype(str).str.strip().str.lower() != "nan"]
            data["Source_File"] = fpath.name
            data["Source_Sheet"] = sheet
            rows.append(data)
    except Exception:
        pass
    return rows


def parse_csv_file(fpath):
    rows = []
    try:
        df_header = pd.read_csv(fpath, nrows=0)
        header_cols = [c.strip().lower() for c in df_header.columns]
        if set(COLUMNS).issubset(set(header_cols)) or "Disco" in df_header.columns:
            df = pd.read_csv(fpath)
            col_map = {}
            for c in COLUMNS:
                for dc in df.columns:
                    if dc.strip().lower() == c.lower() or dc.strip() == c:
                        col_map[dc] = c
                        break
            df = df.rename(columns=col_map)
            for c in COLUMNS:
                if c not in df.columns:
                    df[c] = None
            data = df[COLUMNS].copy()
        else:
            df = pd.read_csv(fpath, header=None)
            header_row = None
            for idx, row in df.iterrows():
                vals = [str(v).strip().lower() for v in row.values[:5]]
                if "disco" in vals and "region" in vals:
                    header_row = idx
                    break
            if header_row is not None:
                data = df.iloc[header_row + 1:].copy()
            else:
                data = df.copy()
            ncols = data.shape[1]
            if ncols > 20:
                data = data.iloc[:, :20]
            elif ncols < 20:
                for i in range(20 - ncols):
                    data[f"_pad_{i}"] = None
            data.columns = SOURCE_COLUMNS

        data = data.dropna(subset=["Disco", "Region"], how="all")
        data = data[~data["Disco"].astype(str).str.contains("NO TCN ATTRIBUTION", case=False, na=False)]
        data = data[data["Disco"].notna() & (data["Disco"].astype(str).str.strip() != "")]
        data = data[data["Disco"].astype(str).str.strip().str.lower() != "nan"]
        data["Source_File"] = fpath.name
        data["Source_Sheet"] = "CSV"
        rows.append(data)
    except Exception:
        pass
    return rows


def parse_word_file(fpath):
    from docx import Document
    rows = []
    try:
        doc = Document(fpath)
        for table in doc.tables:
            header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            if "disco" in header_cells and "region" in header_cells:
                table_data = []
                for row in table.rows[1:]:
                    table_data.append([c.text.strip() for c in row.cells])
                df = pd.DataFrame(table_data)
                ncols = df.shape[1]
                if ncols > 20:
                    df = df.iloc[:, :20]
                elif ncols < 20:
                    for i in range(20 - ncols):
                        df[f"_pad_{i}"] = None
                df.columns = COLUMNS
                df = df.dropna(subset=["Disco", "Region"], how="all")
                df = df[df["Disco"].notna() & (df["Disco"].astype(str).str.strip() != "")]
                df = df[df["Disco"].astype(str).str.strip().str.lower() != "nan"]
                df["Source_File"] = fpath.name
                df["Source_Sheet"] = "Word Table"
                rows.append(df)
    except Exception:
        pass
    return rows


def clean_combined(combined):
    if "Party_Responsible" in combined.columns:
        combined = combined.drop(columns=["Party_Responsible"])
    combined["Date_Off"] = pd.to_datetime(combined["Date_Off"], errors="coerce", dayfirst=True)
    combined["Date_On"] = pd.to_datetime(combined["Date_On"], errors="coerce", dayfirst=True)
    combined["Last_Load_MW"] = pd.to_numeric(combined["Last_Load_MW"], errors="coerce")
    if "Frequency_Hz" not in combined.columns:
        combined["Frequency_Hz"] = None
    combined["Frequency_Hz"] = pd.to_numeric(combined["Frequency_Hz"], errors="coerce")

    def parse_duration(val):
        if pd.isna(val):
            return 0.0
        s = str(val).strip()
        m = re.match(r"(\d+):(\d+):(\d+)", s)
        if m:
            return int(m.group(1)) + int(m.group(2)) / 60 + int(m.group(3)) / 3600
        m = re.match(r"(\d+):(\d+)", s)
        if m:
            return int(m.group(1)) + int(m.group(2)) / 60
        return 0.0

    combined["Duration_Hours"] = combined["Duration"].apply(parse_duration)
    combined["Region"] = combined["Region"].astype(str).str.strip().str.title()
    combined["Disco"] = combined["Disco"].astype(str).str.strip().str.title()
    combined["Class"] = combined["Class"].astype(str).str.strip().str.title()
    combined["Weather_Condition"] = combined["Weather_Condition"].astype(str).str.strip().str.title()
    return combined


@st.cache_data
def load_all_data():
    xlsx_files = sorted(
        list(DATA_DIR.glob("TCN*attributes*.xlsx"))
        + list(DATA_DIR.glob("TCN*attributes*.xlsm"))
        + list(UPLOADS_DIR.glob("*.xlsx"))
        + list(UPLOADS_DIR.glob("*.xlsm")),
        key=lambda f: f.name,
    )
    csv_files = sorted(UPLOADS_DIR.glob("*.csv"), key=lambda f: f.name)
    word_files = sorted(UPLOADS_DIR.glob("*.docx"), key=lambda f: f.name)

    all_rows = []
    for fpath in xlsx_files:
        all_rows.extend(parse_excel_file(fpath))
    for fpath in csv_files:
        all_rows.extend(parse_csv_file(fpath))
    for fpath in word_files:
        all_rows.extend(parse_word_file(fpath))

    if not all_rows:
        return pd.DataFrame(columns=COLUMNS + ["Source_File", "Source_Sheet"])

    combined = pd.concat(all_rows, ignore_index=True)
    return clean_combined(combined)


def apply_filters(df):
    st.sidebar.header("Filters")

    date_min = df["Date_Off"].min()
    date_max = df["Date_Off"].max()
    if pd.notna(date_min) and pd.notna(date_max):
        date_range = st.sidebar.date_input(
            "Date Range",
            value=(date_min.date(), date_max.date()),
            min_value=date_min.date(),
            max_value=date_max.date(),
        )
        if len(date_range) == 2:
            df = df[
                (df["Date_Off"].dt.date >= date_range[0])
                & (df["Date_Off"].dt.date <= date_range[1])
            ]

    regions = sorted(df["Region"].dropna().unique())
    sel_regions = st.sidebar.multiselect("Region", regions, default=regions)
    if sel_regions:
        df = df[df["Region"].isin(sel_regions)]

    discos = sorted(df["Disco"].dropna().unique())
    sel_discos = st.sidebar.multiselect("Disco", discos, default=discos)
    if sel_discos:
        df = df[df["Disco"].isin(sel_discos)]

    classes = sorted(df["Class"].dropna().unique())
    sel_classes = st.sidebar.multiselect("Outage Class", classes, default=classes)
    if sel_classes:
        df = df[df["Class"].isin(sel_classes)]

    search = st.sidebar.text_input("Search (Substation, Feeder, Remarks)")
    if search:
        mask = (
            df["Substation"].astype(str).str.contains(search, case=False, na=False)
            | df["Feeder_33kV"].astype(str).str.contains(search, case=False, na=False)
            | df["Remarks"].astype(str).str.contains(search, case=False, na=False)
        )
        df = df[mask]

    return df


TCN_CHART_LAYOUT = dict(
    font=dict(family="-apple-system, BlinkMacSystemFont, Segoe UI, system-ui, sans-serif", color="#1e2536"),
    paper_bgcolor="rgba(0,0,0,0)",
    plot_bgcolor="rgba(0,0,0,0)",
    title_font=dict(size=14, color="#1e2536", family="-apple-system, BlinkMacSystemFont, Segoe UI, system-ui, sans-serif", weight=600),
    margin=dict(l=40, r=20, t=44, b=40),
    hoverlabel=dict(
        bgcolor="#1e2f5c",
        font_size=12,
        font_family="-apple-system, BlinkMacSystemFont, Segoe UI, system-ui, sans-serif",
        font_color="#eaecf2",
        bordercolor="#1e2f5c",
    ),
)
TCN_COLORS = ["#1e3a7a", "#b8232e", "#3460a8", "#d44450", "#5080c0", "#e87080", "#142a5c", "#901828"]
TCN_RED = "#c81e28"
TCN_RED_SCALE = [[0, "#e8edf5"], [0.5, TCN_RED], [1, "#8b0000"]]


def show_dashboard(df):
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Total Outages", f"{len(df):,}")
    c2.metric("Total Duration (hrs)", f"{df['Duration_Hours'].sum():,.1f}")
    c3.metric("Avg Duration (hrs)", f"{df['Duration_Hours'].mean():,.2f}" if len(df) else "0")
    c4.metric("Total Load Lost (MW)", f"{df['Last_Load_MW'].sum():,.1f}")
    c5.metric("Unique Substations", f"{df['Substation'].nunique():,}")

    col1, col2 = st.columns(2)

    with col1:
        by_region = df.groupby("Region").agg(
            Outages=("Region", "size"),
            Total_Duration=("Duration_Hours", "sum"),
            Total_Load_MW=("Last_Load_MW", "sum"),
        ).reset_index().sort_values("Outages", ascending=False)
        fig = px.bar(
            by_region, x="Region", y="Outages", color="Total_Load_MW",
            title="Outages by Region",
            color_continuous_scale=TCN_RED_SCALE,
            labels={"Total_Load_MW": "Load Lost (MW)"},
        )
        fig.update_layout(height=400, **TCN_CHART_LAYOUT)
        fig.update_traces(marker_line_width=0)
        st.plotly_chart(fig, use_container_width=True)

    with col2:
        by_class = df["Class"].value_counts().reset_index()
        by_class.columns = ["Class", "Count"]
        fig = px.pie(by_class, names="Class", values="Count", title="Outage Classification",
                     hole=0.4, color_discrete_sequence=TCN_COLORS)
        fig.update_layout(height=400, **TCN_CHART_LAYOUT)
        st.plotly_chart(fig, use_container_width=True)

    col3, col4 = st.columns(2)

    with col3:
        if df["Date_Off"].notna().any():
            daily = df.groupby(df["Date_Off"].dt.date).agg(
                Outages=("Date_Off", "size"),
                Total_Load_MW=("Last_Load_MW", "sum"),
            ).reset_index()
            daily.columns = ["Date", "Outages", "Total_Load_MW"]
            fig = go.Figure()
            fig.add_trace(go.Bar(x=daily["Date"], y=daily["Outages"], name="Outages",
                                  yaxis="y", marker_color=TCN_COLORS[0], marker_line_width=0))
            fig.add_trace(go.Scatter(x=daily["Date"], y=daily["Total_Load_MW"], name="Load Lost (MW)",
                                     yaxis="y2", mode="lines+markers",
                                     line=dict(color=TCN_RED, width=2),
                                     marker=dict(size=5, color=TCN_RED)))
            fig.update_layout(
                title="Daily Outages & Load Lost",
                yaxis=dict(title="Outages"),
                yaxis2=dict(title="Load Lost (MW)", overlaying="y", side="right"),
                height=400,
                legend=dict(x=0, y=1.1, orientation="h"),
                **TCN_CHART_LAYOUT,
            )
            st.plotly_chart(fig, use_container_width=True)

    with col4:
        by_event = df["Event_Indication"].astype(str).value_counts().head(10).reset_index()
        by_event.columns = ["Event", "Count"]
        fig = px.bar(by_event, x="Count", y="Event", orientation="h", title="Top 10 Event Indications",
                     color_discrete_sequence=[TCN_COLORS[0]])
        fig.update_layout(height=400, yaxis=dict(autorange="reversed"), **TCN_CHART_LAYOUT)
        fig.update_traces(marker_line_width=0)
        st.plotly_chart(fig, use_container_width=True)

    col5, col6 = st.columns(2)

    with col5:
        by_weather = df["Weather_Condition"].value_counts().reset_index()
        by_weather.columns = ["Weather", "Count"]
        fig = px.pie(by_weather, names="Weather", values="Count", title="Weather During Outages",
                     hole=0.4, color_discrete_sequence=TCN_COLORS)
        fig.update_layout(height=400, **TCN_CHART_LAYOUT)
        st.plotly_chart(fig, use_container_width=True)

    with col6:
        region_duration = df.groupby("Region")["Duration_Hours"].sum().reset_index().sort_values("Duration_Hours", ascending=True)
        fig = px.bar(region_duration, x="Duration_Hours", y="Region", orientation="h",
                     title="Total Outage Duration by Region (hrs)",
                     color_discrete_sequence=[TCN_RED])
        fig.update_layout(height=400, **TCN_CHART_LAYOUT)
        fig.update_traces(marker_line_width=0)
        st.plotly_chart(fig, use_container_width=True)


def show_data_table(df):
    st.header("Outage Records")
    display = df.copy()
    display_cols = {old: new for old, new in zip(COLUMNS, DISPLAY_COLUMNS) if old in display.columns}
    display = display.rename(columns=display_cols)
    show_cols = DISPLAY_COLUMNS + ["Source_File", "Source_Sheet"]
    show_cols = [c for c in show_cols if c in display.columns]
    st.dataframe(display[show_cols], use_container_width=True, height=500)
    st.caption(f"Showing {len(display):,} records")


def show_region_analysis(df):
    st.header("Region Analysis")
    region = st.selectbox("Select Region", sorted(df["Region"].dropna().unique()))
    rdf = df[df["Region"] == region]

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Outages", f"{len(rdf):,}")
    c2.metric("Total Duration (hrs)", f"{rdf['Duration_Hours'].sum():,.1f}")
    c3.metric("Total Load Lost (MW)", f"{rdf['Last_Load_MW'].sum():,.1f}")
    c4.metric("Substations Affected", f"{rdf['Substation'].nunique()}")

    col1, col2 = st.columns(2)
    with col1:
        by_sub = rdf.groupby("Substation").agg(
            Outages=("Substation", "size"),
            Load_MW=("Last_Load_MW", "sum"),
        ).reset_index().sort_values("Outages", ascending=False).head(15)
        fig = px.bar(by_sub, x="Substation", y="Outages", color="Load_MW",
                     title=f"Top Substations: {region}",
                     color_continuous_scale=TCN_RED_SCALE)
        fig.update_layout(height=400, xaxis_tickangle=-45, **TCN_CHART_LAYOUT)
        fig.update_traces(marker_line_width=0)
        st.plotly_chart(fig, use_container_width=True)

    with col2:
        by_feeder = rdf.groupby("Feeder_33kV").agg(
            Outages=("Feeder_33kV", "size"),
            Duration=("Duration_Hours", "sum"),
        ).reset_index().sort_values("Duration", ascending=False).head(15)
        fig = px.bar(by_feeder, x="Duration", y="Feeder_33kV", orientation="h",
                     title=f"Top Feeders by Duration: {region}",
                     color_discrete_sequence=[TCN_COLORS[0]])
        fig.update_layout(height=400, **TCN_CHART_LAYOUT)
        fig.update_traces(marker_line_width=0)
        st.plotly_chart(fig, use_container_width=True)

    st.subheader(f"Records: {region}")
    display = rdf.copy()
    display_cols = {old: new for old, new in zip(COLUMNS, DISPLAY_COLUMNS) if old in display.columns}
    display = display.rename(columns=display_cols)
    show_cols = [c for c in DISPLAY_COLUMNS if c in display.columns]
    st.dataframe(display[show_cols], use_container_width=True, height=400)


def show_export(df):
    st.header("Export Reports")

    export_type = st.radio("Export Format", ["Excel (.xlsx)", "CSV (.csv)"], horizontal=True)

    if st.button("Generate Export"):
        if export_type == "CSV (.csv)":
            display = df.copy()
            display_cols = {old: new for old, new in zip(COLUMNS, DISPLAY_COLUMNS) if old in display.columns}
            display = display.rename(columns=display_cols)
            csv = display.to_csv(index=False)
            st.download_button("Download CSV", csv, "tcn_outage_report.csv", "text/csv")
        else:
            buf = io.BytesIO()
            display = df.copy()
            display_cols = {old: new for old, new in zip(COLUMNS, DISPLAY_COLUMNS) if old in display.columns}
            display = display.rename(columns=display_cols)

            with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
                display.to_excel(writer, sheet_name="All Records", index=False)

                summary = df.groupby("Region").agg(
                    Total_Outages=("Region", "size"),
                    Total_Duration_Hrs=("Duration_Hours", "sum"),
                    Avg_Duration_Hrs=("Duration_Hours", "mean"),
                    Total_Load_Lost_MW=("Last_Load_MW", "sum"),
                    Unique_Substations=("Substation", "nunique"),
                ).reset_index()
                summary.to_excel(writer, sheet_name="Summary by Region", index=False)

                if df["Date_Off"].notna().any():
                    daily = df.groupby(df["Date_Off"].dt.date).agg(
                        Total_Outages=("Date_Off", "size"),
                        Total_Duration_Hrs=("Duration_Hours", "sum"),
                        Total_Load_Lost_MW=("Last_Load_MW", "sum"),
                    ).reset_index()
                    daily.columns = ["Date", "Total_Outages", "Total_Duration_Hrs", "Total_Load_Lost_MW"]
                    daily.to_excel(writer, sheet_name="Daily Summary", index=False)

                for ws_name in writer.sheets:
                    ws = writer.sheets[ws_name]
                    ws.set_column(0, 20, 18)

            st.download_button("Download Excel", buf.getvalue(), "tcn_outage_report.xlsx",
                               "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    st.subheader("Summary Statistics")
    summary = df.groupby("Region").agg(
        Total_Outages=("Region", "size"),
        Total_Duration_Hrs=("Duration_Hours", "sum"),
        Avg_Duration_Hrs=("Duration_Hours", "mean"),
        Total_Load_Lost_MW=("Last_Load_MW", "sum"),
        Unique_Substations=("Substation", "nunique"),
    ).reset_index().sort_values("Total_Outages", ascending=False)
    summary["Total_Duration_Hrs"] = summary["Total_Duration_Hrs"].round(1)
    summary["Avg_Duration_Hrs"] = summary["Avg_Duration_Hrs"].round(2)
    summary["Total_Load_Lost_MW"] = summary["Total_Load_Lost_MW"].round(1)
    st.dataframe(summary, use_container_width=True)


def show_outage_entry():
    st.header("Report New Outage")
    st.markdown("Select location and equipment using the dropdowns, then fill in the outage details below.")

    # --- Cascading selectors OUTSIDE the form so they trigger reruns ---
    st.subheader("Location")
    loc1, loc2 = st.columns(2)
    region = loc1.selectbox("Region", options=scoped_regions(sorted(SUBSTATION_EQUIPMENT.keys())), key="entry_region")

    subregions = sorted(SUBSTATION_EQUIPMENT.get(region, {}).keys())
    sub_acc = loc2.selectbox("SubRegion / ACC", options=subregions, key="entry_sub")

    loc3, loc4 = st.columns(2)
    substations = get_substations_for(region, sub_acc)
    substation = loc3.selectbox("Substation", options=substations, key="entry_station")

    suggested_discos = REGION_DISCO_MAP.get(region, DISCOS)
    all_discos = suggested_discos + [d for d in DISCOS if d not in suggested_discos]
    disco = loc4.selectbox("Disco", options=all_discos, key="entry_disco")

    st.divider()
    st.subheader("Equipment")
    eq1, eq2 = st.columns(2)
    equip_type = eq1.selectbox("Equipment Type", options=EQUIPMENT_TYPES, key="entry_equip_type")

    equip_list = get_equipment_list(region, sub_acc, substation, equip_type)
    if equip_list:
        equipment = eq2.selectbox("Equipment", options=equip_list, key="entry_equip")
    else:
        equipment = eq2.text_input("Equipment (specify)", key="entry_equip_text")

    feeder = st.text_input("33kV Feeder Name (if applicable)")

    # --- Remaining fields inside the form for batch submission ---
    st.divider()
    with st.form("outage_entry", clear_on_submit=True):
        st.subheader("Outage Timing")
        t1, t2, t3, t4 = st.columns(4)
        date_off = t1.date_input("Date Off")
        time_off = t2.time_input("Time Off")
        date_on = t3.date_input("Date On")
        time_on = t4.time_input("Time On")

        t5, t6 = st.columns(2)
        duration = t5.text_input("Duration (H:mm)", placeholder="01:30")
        outage_class = t6.selectbox("Outage Class", options=OUTAGE_CLASSES)

        st.divider()
        st.subheader("Event Details")
        e1, e2, e3 = st.columns(3)
        load_mw = e1.number_input("Last Load (MW)", min_value=0.0, step=0.1)
        frequency = e2.number_input("Frequency (Hz)", min_value=0.0, max_value=100.0, value=50.0, step=0.01)
        event = e3.selectbox("Event / Indication", options=EVENT_INDICATIONS)

        e4, e5, e6 = st.columns(3)
        officer_int = e4.text_input("Officer (Interruption)")
        officer_res = e5.text_input("Officer (Restoration)")
        weather = e6.selectbox("Weather Condition", options=WEATHER_CONDITIONS)

        remarks = st.text_area("Remarks")

        submitted = st.form_submit_button("Submit Outage Report", type="primary")
        if submitted:
            if not substation or not region:
                st.error("Region and Substation are required.")
            else:
                equip_name = equipment if equipment else ""
                equip_full = f"[{equip_type}] {equip_name}" if equip_name else equip_type
                new_row = pd.DataFrame([{
                    "Disco": disco,
                    "Region": region,
                    "SubRegion_ACC": sub_acc,
                    "Substation": substation,
                    "Feeder_33kV": feeder,
                    "Equipment_Type": equip_type,
                    "Equipment": equip_name,
                    "Date_Off": date_off.strftime("%d/%m/%Y"),
                    "Hour_Off": f"{time_off.hour:02d}:00:00",
                    "Minute_Off": f"00:{time_off.minute:02d}:00",
                    "Date_On": date_on.strftime("%d/%m/%Y"),
                    "Hour_On": f"{time_on.hour:02d}:00:00",
                    "Minute_On": f"00:{time_on.minute:02d}:00",
                    "Duration": duration or "00:00",
                    "Class": outage_class,
                    "Last_Load_MW": load_mw,
                    "Frequency_Hz": frequency,
                    "Event_Indication": event,
                    "Officer_Interruption": officer_int,
                    "Officer_Restoration": officer_res,
                    "Weather_Condition": weather,
                    "Remarks": remarks,
                }])
                manual_file = UPLOADS_DIR / "manual_entries.csv"
                if manual_file.exists():
                    existing_df = pd.read_csv(manual_file)
                    combined = pd.concat([existing_df, new_row], ignore_index=True)
                else:
                    combined = new_row
                combined.to_csv(manual_file, index=False)
                load_all_data.clear()
                st.success(f"Outage report submitted: {equip_full} at {substation} ({region} / {sub_acc})")
                log_activity("manual_outage_entry", f"{equip_full} at {substation} ({region} / {sub_acc})")
                st.rerun()


def show_upload():
    st.header("Upload Files")
    st.markdown(
        "Upload **Excel** (.xlsx/.xlsm), **CSV** (.csv), or **Word** (.docx) files. "
        "Files with the standard TCN outage columns (Disco, Region, Substation, etc.) "
        "will be parsed automatically."
    )

    uploaded_files = st.file_uploader(
        "Choose files to upload",
        type=["xlsx", "xlsm", "csv", "docx"],
        accept_multiple_files=True,
    )

    if uploaded_files:
        for uf in uploaded_files:
            dest = UPLOADS_DIR / uf.name
            st.markdown(f"**{uf.name}** ({uf.size / 1024:.1f} KB)")

            suffix = Path(uf.name).suffix.lower()
            preview_rows = []

            if suffix in (".xlsx", ".xlsm"):
                tmp = UPLOADS_DIR / f"_tmp_{uf.name}"
                tmp.write_bytes(uf.getvalue())
                preview_rows = parse_excel_file(tmp)
                tmp.unlink(missing_ok=True)
            elif suffix == ".csv":
                tmp = UPLOADS_DIR / f"_tmp_{uf.name}"
                tmp.write_bytes(uf.getvalue())
                preview_rows = parse_csv_file(tmp)
                tmp.unlink(missing_ok=True)
            elif suffix == ".docx":
                tmp = UPLOADS_DIR / f"_tmp_{uf.name}"
                tmp.write_bytes(uf.getvalue())
                preview_rows = parse_word_file(tmp)
                if not preview_rows:
                    from docx import Document as DocxDocument
                    try:
                        doc = DocxDocument(tmp)
                        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                        if text:
                            with st.expander(f"Text content from {uf.name}"):
                                st.text_area("Document text", text, height=300, disabled=True)
                    except Exception:
                        pass
                tmp.unlink(missing_ok=True)

            if preview_rows:
                preview_df = pd.concat(preview_rows, ignore_index=True)
                preview_display = preview_df.copy()
                col_map = {old: new for old, new in zip(COLUMNS, DISPLAY_COLUMNS) if old in preview_display.columns}
                preview_display = preview_display.rename(columns=col_map)
                show = [c for c in DISPLAY_COLUMNS if c in preview_display.columns]
                st.success(f"Found **{len(preview_df)}** outage records")
                st.dataframe(preview_display[show], use_container_width=True, height=250)
            elif suffix != ".docx":
                st.warning(f"No outage records found in {uf.name}.")

        if st.button("Save Uploaded Files", type="primary"):
            saved = []
            upload_region = None if is_super_admin() else current_region()
            for uf in uploaded_files:
                dest = UPLOADS_DIR / uf.name
                dest.write_bytes(uf.getvalue())
                saved.append(uf.name)
                # exempt_from_expiry=True: this folder is also this page's live
                # data source (load_all_data() reads it back), so these files
                # must NOT be swept by the 90-day purge like other uploads are.
                save_uploaded_file(uf, "TCN Outage Manager", upload_region, exempt_from_expiry=True)
            load_all_data.clear()
            st.success(f"Saved {len(saved)} file(s): {', '.join(saved)}")
            log_activity("tcn_file_upload", f"Saved {len(saved)} file(s): {', '.join(saved)}")
            st.rerun()

    st.divider()
    st.subheader("Manage Uploaded Files")
    existing = sorted(UPLOADS_DIR.glob("*"))
    existing = [f for f in existing if f.is_file() and not f.name.startswith("_tmp_")]
    if existing:
        for f in existing:
            col1, col2, col3 = st.columns([4, 1, 1])
            col1.text(f.name)
            col2.text(f"{f.stat().st_size / 1024:.1f} KB")
            if col3.button("Remove", key=f"rm_{f.name}"):
                f.unlink()
                load_all_data.clear()
                st.rerun()
    else:
        st.info("No uploaded files yet.")


def show_hierarchy_browser():
    st.header("TCN Network Hierarchy")
    st.markdown("Browse all regions, sub-regions/ACCs, substations, and their equipment.")

    view_mode = st.radio("View", ["By Region", "Full Equipment Table"], horizontal=True)

    if view_mode == "By Region":
        region = st.selectbox("Select Region", sorted(SUBSTATION_EQUIPMENT.keys()), key="hier_region")
        subregions = SUBSTATION_EQUIPMENT[region]
        for sub_name in sorted(subregions.keys()):
            stations = subregions[sub_name]
            with st.expander(f"{sub_name} ({len(stations)} substations)", expanded=True):
                for s_name in sorted(stations.keys()):
                    equip = stations[s_name]
                    if not equip:
                        st.markdown(f"**{s_name}**")
                        continue
                    total_equip = sum(len(v) for v in equip.values())
                    st.markdown(f"**{s_name}** ({total_equip} equipment)")
                    for etype in ["Transmission Line", "Transformer", "Reactor", "Capacitor Bank"]:
                        items = equip.get(etype, [])
                        if items:
                            for item in items:
                                st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;`{etype[:5]}` {item}")
    else:
        rows = []
        for region, subs in sorted(SUBSTATION_EQUIPMENT.items()):
            for sub, stations in sorted(subs.items()):
                for station, equip in sorted(stations.items()):
                    if not equip:
                        rows.append({"Region": region, "SubRegion/ACC": sub, "Substation": station, "Type": "", "Equipment": ""})
                    for etype, items in equip.items():
                        for item in items:
                            rows.append({"Region": region, "SubRegion/ACC": sub, "Substation": station, "Type": etype, "Equipment": item})
        hier_df = pd.DataFrame(rows)
        st.dataframe(hier_df, use_container_width=True, height=600)
        total_equip = len([r for r in rows if r["Equipment"]])
        total_stations = len(set((r["Region"], r["SubRegion/ACC"], r["Substation"]) for r in rows))
        st.caption(f"{total_equip} equipment items across {total_stations} substations in {len(SUBSTATION_EQUIPMENT)} regions")


def show_generate_report(df):
    st.header("Generate Outage Report")

    mode = st.radio("Report Period", ["Today", "Yesterday", "Custom Days", "Date Range"], horizontal=True)

    today = date.today()
    if mode == "Today":
        start_date = today
        end_date = today
        st.info(f"Report for: **{today.strftime('%d %B %Y')}**")
    elif mode == "Yesterday":
        start_date = today - timedelta(days=1)
        end_date = start_date
        st.info(f"Report for: **{start_date.strftime('%d %B %Y')}**")
    elif mode == "Custom Days":
        num_days = st.number_input("Number of days (from today going back)", min_value=1, max_value=365, value=7)
        start_date = today - timedelta(days=num_days - 1)
        end_date = today
        st.info(f"Report period: **{start_date.strftime('%d %B %Y')}** to **{end_date.strftime('%d %B %Y')}** ({num_days} days)")
    else:
        dr = st.date_input("Select Date Range", value=(today - timedelta(days=6), today))
        if isinstance(dr, tuple) and len(dr) == 2:
            start_date, end_date = dr
        else:
            start_date = end_date = today

    rdf = df[df["Date_Off"].notna()].copy()
    rdf = rdf[(rdf["Date_Off"].dt.date >= start_date) & (rdf["Date_Off"].dt.date <= end_date)]

    st.divider()
    if rdf.empty:
        st.warning("No outage records found for the selected period.")
        return

    st.subheader(f"Report Summary ({len(rdf)} outages)")
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Total Outages", f"{len(rdf):,}")
    m2.metric("Total Duration (hrs)", f"{rdf['Duration_Hours'].sum():,.1f}")
    m3.metric("Total Load Lost (MW)", f"{rdf['Last_Load_MW'].sum():,.1f}")
    m4.metric("Substations Affected", f"{rdf['Substation'].nunique()}")

    by_region = rdf.groupby("Region").agg(
        Outages=("Region", "size"),
        Duration_Hrs=("Duration_Hours", "sum"),
        Load_Lost_MW=("Last_Load_MW", "sum"),
        Substations=("Substation", "nunique"),
    ).reset_index().sort_values("Outages", ascending=False)
    by_region["Duration_Hrs"] = by_region["Duration_Hrs"].round(1)
    by_region["Load_Lost_MW"] = by_region["Load_Lost_MW"].round(1)
    st.dataframe(by_region, use_container_width=True)

    col1, col2 = st.columns(2)
    with col1:
        fig = px.bar(by_region, x="Region", y="Outages", color="Load_Lost_MW",
                     title="Outages by Region",
                     color_continuous_scale=TCN_RED_SCALE)
        fig.update_layout(height=350, **TCN_CHART_LAYOUT)
        fig.update_traces(marker_line_width=0)
        st.plotly_chart(fig, use_container_width=True)
    with col2:
        by_class = rdf["Class"].value_counts().reset_index()
        by_class.columns = ["Class", "Count"]
        fig = px.pie(by_class, names="Class", values="Count", title="Outage Classification",
                     hole=0.4, color_discrete_sequence=TCN_COLORS)
        fig.update_layout(height=350, **TCN_CHART_LAYOUT)
        st.plotly_chart(fig, use_container_width=True)

    st.divider()
    st.subheader("Download Report")
    report_name = f"TCN_Outage_Report_{start_date.strftime('%d-%m-%Y')}_to_{end_date.strftime('%d-%m-%Y')}"

    dl1, dl2 = st.columns(2)

    with dl1:
        buf = io.BytesIO()
        display = rdf.copy()
        display_cols = {old: new for old, new in zip(COLUMNS, DISPLAY_COLUMNS) if old in display.columns}
        display = display.rename(columns=display_cols)

        with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
            summary_data = pd.DataFrame([{
                "Report Period": f"{start_date.strftime('%d/%m/%Y')} to {end_date.strftime('%d/%m/%Y')}",
                "Generated": datetime.now().strftime("%d/%m/%Y %H:%M"),
                "Generated By": st.session_state.get("username", ""),
                "Total Outages": len(rdf),
                "Total Duration (hrs)": round(rdf["Duration_Hours"].sum(), 1),
                "Total Load Lost (MW)": round(rdf["Last_Load_MW"].sum(), 1),
                "Substations Affected": rdf["Substation"].nunique(),
            }])
            summary_data.T.reset_index().rename(columns={"index": "Metric", 0: "Value"}).to_excel(
                writer, sheet_name="Report Info", index=False
            )

            show_cols = [c for c in DISPLAY_COLUMNS if c in display.columns]
            display[show_cols].to_excel(writer, sheet_name="Outage Records", index=False)

            by_region.to_excel(writer, sheet_name="By Region", index=False)

            by_station = rdf.groupby(["Region", "Substation"]).agg(
                Outages=("Substation", "size"),
                Duration_Hrs=("Duration_Hours", "sum"),
                Load_Lost_MW=("Last_Load_MW", "sum"),
            ).reset_index().sort_values("Outages", ascending=False)
            by_station["Duration_Hrs"] = by_station["Duration_Hrs"].round(1)
            by_station["Load_Lost_MW"] = by_station["Load_Lost_MW"].round(1)
            by_station.to_excel(writer, sheet_name="By Substation", index=False)

            if rdf["Date_Off"].notna().any():
                daily = rdf.groupby(rdf["Date_Off"].dt.date).agg(
                    Outages=("Date_Off", "size"),
                    Duration_Hrs=("Duration_Hours", "sum"),
                    Load_Lost_MW=("Last_Load_MW", "sum"),
                ).reset_index()
                daily.columns = ["Date", "Outages", "Duration_Hrs", "Load_Lost_MW"]
                daily.to_excel(writer, sheet_name="Daily Breakdown", index=False)

            for ws_name in writer.sheets:
                ws = writer.sheets[ws_name]
                ws.set_column(0, 20, 18)

        st.download_button(
            "Download Excel Report",
            buf.getvalue(),
            f"{report_name}.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            type="primary",
            use_container_width=True,
        )

    with dl2:
        display2 = rdf.copy()
        display_cols2 = {old: new for old, new in zip(COLUMNS, DISPLAY_COLUMNS) if old in display2.columns}
        display2 = display2.rename(columns=display_cols2)
        show_cols2 = [c for c in DISPLAY_COLUMNS if c in display2.columns]
        csv = display2[show_cols2].to_csv(index=False)
        st.download_button(
            "Download CSV Report",
            csv,
            f"{report_name}.csv",
            "text/csv",
            use_container_width=True,
        )

    st.divider()
    st.subheader("Outage Records")
    display3 = rdf.copy()
    display_cols3 = {old: new for old, new in zip(COLUMNS, DISPLAY_COLUMNS) if old in display3.columns}
    display3 = display3.rename(columns=display_cols3)
    show_cols3 = [c for c in DISPLAY_COLUMNS if c in display3.columns]
    st.dataframe(display3[show_cols3], use_container_width=True, height=400)


def main():
    st.markdown(f"""
    <div style="display: flex; align-items: center; gap: 0.75rem; margin-bottom: 0.5rem;">
        <img src="data:image/png;base64,{_logo_b64}" alt="TCN"
             style="height: 44px; width: auto; flex-shrink: 0;
                    filter: drop-shadow(0 1px 2px oklch(32% 0.06 250 / 0.08));" />
        <div>
            <h1 style="margin: 0; color: oklch(32% 0.12 250); font-weight: 700; font-size: 1.5rem;
                       letter-spacing: -0.02em; line-height: 1.2;
                       font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', system-ui, sans-serif;">
                TCN Outage Manager
            </h1>
            <p style="color: oklch(45% 0.006 250); font-size: 0.8rem; margin: 0.125rem 0 0;
                      letter-spacing: 0.01em;">
                Outage Tracking and Reporting (file-based attribute records)
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    df = load_all_data()
    df = filter_to_user_region(df, region_col="Region")

    tab_names = [
        "Dashboard", "Records", "Region Analysis", "Generate Report",
        "Report Outage", "Network Hierarchy", "Export", "Upload Files",
    ]

    tabs = st.tabs(tab_names)
    tab_map = {name: tab for name, tab in zip(tab_names, tabs)}

    if df.empty:
        with tab_map["Report Outage"]:
            show_outage_entry()
        with tab_map["Network Hierarchy"]:
            show_hierarchy_browser()
        with tab_map["Upload Files"]:
            show_upload()
        with tab_map["Dashboard"]:
            st.warning("No outage data found. Report an outage or upload TCN attribute files to get started.")
        return

    filtered = apply_filters(df)

    with tab_map["Dashboard"]:
        show_dashboard(filtered)
    with tab_map["Records"]:
        show_data_table(filtered)
    with tab_map["Region Analysis"]:
        show_region_analysis(filtered)
    with tab_map["Generate Report"]:
        show_generate_report(filtered)
    with tab_map["Report Outage"]:
        show_outage_entry()
    with tab_map["Network Hierarchy"]:
        show_hierarchy_browser()
    with tab_map["Export"]:
        show_export(filtered)
    with tab_map["Upload Files"]:
        show_upload()


if __name__ == "__main__":
    main()
