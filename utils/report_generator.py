"""
Report Generator Utility
Generates comprehensive outage reports in both Word (.docx) and PDF formats.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from typing import List, Tuple
from datetime import datetime
import os
from PIL import Image as PILImage
import io


def add_table_to_word(doc: Document, df: pd.DataFrame, title: str = None):
    """Add a pandas DataFrame as a table to a Word document."""
    if title:
        # Add title
        heading = doc.add_paragraph(title, style='Heading 2')
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if df.empty:
        return
    
    # Format numerical columns to 2 decimal places
    df_display = df.copy()
    for col in df_display.columns:
        if df_display[col].dtype in ['float64', 'float32', 'int64', 'int32']:
            df_display[col] = df_display[col].apply(lambda x: f"{float(x):.2f}" if pd.notna(x) else '')
    
    # Reset index to ensure sequential row numbers
    df_display = df_display.reset_index(drop=True)
    
    # Create table with headers
    rows, cols = df_display.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_display.columns):
        if i < len(hdr_cells):
            hdr_cells[i].text = str(col)
            # Format header row
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '4472C4')
                hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
    
    # Add data rows - use enumerate to get sequential row numbers
    for row_idx, (_, row) in enumerate(df_display.iterrows(), start=1):
        row_cells = table.rows[row_idx].cells
        for j, val in enumerate(row):
            if j < len(row_cells):
                cell_text = str(val) if pd.notna(val) else ''
                row_cells[j].text = cell_text
    
    doc.add_paragraph()  # Add spacing after table


def generate_word_report(
    region: str,
    start_date: str,
    end_date: str,
    summary_stats: dict,
    outage_df: pd.DataFrame,
    station_summary: pd.DataFrame,
    feeder_summary: pd.DataFrame,
    cause_summary: pd.DataFrame,
    party_summary: pd.DataFrame,
    image_paths: List[str] = None,
    output_path: str = None,
    feeder_party_pivot: pd.DataFrame = None,
    feeder_party_violations: pd.DataFrame = None,
    top_prolonged_outages: pd.DataFrame = None,
    wrong_attribution: pd.DataFrame = None,
    missing_values_outages: pd.DataFrame = None,
    unrestored_feeders: pd.DataFrame = None
) -> str:
    """
    Generate a Word (.docx) report with outage analytics data.
    
    Parameters
    ----------
    region : str
        The region name
    start_date : str
        Start date in format 'YYYY-MM-DD'
    end_date : str
        End date in format 'YYYY-MM-DD'
    summary_stats : dict
        Dictionary with keys: 'num_outages', 'total_minutes', 'avg_duration_min'
    outage_df : pd.DataFrame
        Raw outage records
    station_summary : pd.DataFrame
        Station-level summary (outages_count, total_outage_min, avg_outage_min)
    feeder_summary : pd.DataFrame
        Feeder-level summary
    cause_summary : pd.DataFrame
        Outage cause breakdown
    party_summary : pd.DataFrame
        Party responsible breakdown
    image_paths : List[str], optional
        List of image paths (charts) to embed
    output_path : str, optional
        Output file path. If None, generates default name.
    feeder_party_pivot : pd.DataFrame, optional
        Feeder-Party responsible pivot table
    feeder_party_violations : pd.DataFrame, optional
        SLA violations (negative TCN availability)
    top_prolonged_outages : pd.DataFrame, optional
        Top 20% most prolonged outages
    wrong_attribution : pd.DataFrame, optional
        Outages with wrong party attribution
    missing_values_outages : pd.DataFrame, optional
        Outages with missing values (excluding date_on/time_on)
    unrestored_feeders : pd.DataFrame, optional
        Feeders yet to be restored (no date_on/time_on)
    
    Returns
    -------
    str
        Path to generated Word document
    """
    
    if output_path is None:
        output_path = f"Outage_Report_{region}_{start_date}_{end_date}.docx"
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f'Outage Report - {region}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date range
    date_para = doc.add_paragraph(f'Report Period: {start_date} to {end_date}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Generated date
    gen_date = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    
    # Summary metrics table
    try:
        summary_data = {
            'Metric': ['Number of Outages', 'Total Outage Minutes', 'Total Outage Hours', 'Average Duration (minutes)', 'Total Load Loss (MWh)'],
            'Value': [
                str(summary_stats.get('num_outages', 0)),
                f"{float(summary_stats.get('total_minutes', 0)):.2f}",
                f"{float(summary_stats.get('total_minutes', 0)) / 60:.2f}",
                f"{float(summary_stats.get('avg_duration_min', 0)):.2f}",
                f"{float(summary_stats.get('total_load_loss_mwh', 0)):.2f}"
            ]
        }
        summary_df = pd.DataFrame(summary_data)
        add_table_to_word(doc, summary_df)
    except Exception as e:
        doc.add_paragraph(f"[Error adding Summary table: {str(e)}]")
    
    # Outage Details
    doc.add_heading('Detailed Analysis', 1)
    
    # Cause Analysis
    if not cause_summary.empty:
        try:
            add_table_to_word(doc, cause_summary.head(10), 'Outage Causes (Top 10)')
        except Exception as e:
            doc.add_paragraph(f"[Error adding Outage Causes table: {str(e)}]")
    
    # Party Responsible
    if not party_summary.empty:
        try:
            add_table_to_word(doc, party_summary, 'Party Responsible')
        except Exception as e:
            doc.add_paragraph(f"[Error adding Party Responsible table: {str(e)}]")
    
    # Station Summary
    if not station_summary.empty:
        try:
            add_table_to_word(doc, station_summary.head(15), 'Top Stations by Outage Minutes (Top 15)')
        except Exception as e:
            doc.add_paragraph(f"[Error adding Station Summary table: {str(e)}]")
    
    # Feeder Summary
    if not feeder_summary.empty:
        try:
            display_cols = ['feeder_33kv', 'outages_count', 'outage_hrs', 'avg_outage_hrs', 'total_load_loss_mwh', 'avg_load_loss_mwh']
            feeder_display = feeder_summary[display_cols] if all(c in feeder_summary.columns for c in display_cols) else feeder_summary.head(10)
            add_table_to_word(doc, feeder_display.head(15), 'Top Feeders by Outage Hours (Top 15)')
        except Exception as e:
            doc.add_paragraph(f"[Error adding Feeder Summary table: {str(e)}]")
    
    # Outage Table (Feeder-Party breakdown)
    if feeder_party_pivot is not None and not feeder_party_pivot.empty:
        try:
            add_table_to_word(doc, feeder_party_pivot.head(25), 'Outage Table - Feeder by Party Responsible')
        except Exception as e:
            doc.add_paragraph(f"[Error adding Outage Table: {str(e)}]")
    
    # SLA Violations (Negative TCN Availability)
    if feeder_party_violations is not None and not feeder_party_violations.empty:
        try:
            add_table_to_word(doc, feeder_party_violations, 'SLA Violations - TCN Negative Availability (TCN < 0)')
        except Exception as e:
            doc.add_paragraph(f"[Error adding SLA Violations table: {str(e)}]")
    
    # Top 20% Most Prolonged Outages
    if top_prolonged_outages is not None and not top_prolonged_outages.empty:
        try:
            add_table_to_word(doc, top_prolonged_outages, 'Top 20% Most Prolonged Outages')
        except Exception as e:
            doc.add_paragraph(f"[Error adding Prolonged Outages table: {str(e)}]")
    
    # Wrong Attribution for Party Responsible
    if wrong_attribution is not None and not wrong_attribution.empty:
        try:
            add_table_to_word(doc, wrong_attribution, 'Wrong Attribution for Party Responsible')
        except Exception as e:
            doc.add_paragraph(f"[Error adding Wrong Attribution table: {str(e)}]")
    
    # Outages with Missing Values
    if missing_values_outages is not None and not missing_values_outages.empty:
        try:
            add_table_to_word(doc, missing_values_outages, 'Outages with Missing Values')
        except Exception as e:
            doc.add_paragraph(f"[Error adding Missing Values table: {str(e)}]")
    
    # Feeders Yet to be Restored
    if unrestored_feeders is not None and not unrestored_feeders.empty:
        try:
            add_table_to_word(doc, unrestored_feeders, 'Feeders Yet to be Restored')
        except Exception as e:
            doc.add_paragraph(f"[Error adding Unrestored Feeders table: {str(e)}]")
    
    # Raw Outages Table
    if not outage_df.empty:
        try:
            display_cols = ['date_off', 'time_off', 'date_on', 'time_on', 'station', 'outage_class', 'party_responsible', 'weather_condition', 'load_loss_mwh']
            available_cols = [c for c in display_cols if c in outage_df.columns]
            outage_display = outage_df[available_cols].head(50)
            add_table_to_word(doc, outage_display, 'Recent Outage Records (Latest 50)')
        except Exception as e:
            doc.add_paragraph(f"[Error adding Outage Records table: {str(e)}]")
    
    # Add images/charts if provided
    if image_paths:
        doc.add_page_break()
        doc.add_heading('Charts & Visualizations', 1)
        
        for img_path in image_paths:
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6))
                    doc.add_paragraph()
                except Exception as e:
                    doc.add_paragraph(f"[Could not embed image: {img_path}]")
    
    # Save document
    doc.save(output_path)
    return output_path


def generate_pdf_report_with_tables(
    region: str,
    start_date: str,
    end_date: str,
    summary_stats: dict,
    outage_df: pd.DataFrame,
    station_summary: pd.DataFrame,
    feeder_summary: pd.DataFrame,
    image_paths: List[str] = None,
    output_path: str = None,
    feeder_party_pivot: pd.DataFrame = None,
    feeder_party_violations: pd.DataFrame = None
) -> str:
    """
    Generate a PDF report using ReportLab with tables.
    
    Returns
    -------
    str
        Path to generated PDF
    """
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    
    if output_path is None:
        output_path = f"Outage_Report_{region}_{start_date}_{end_date}.pdf"
    
    # Use landscape for better table display
    doc = SimpleDocTemplate(output_path, pagesize=landscape(A4))
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a3a6b'),
        spaceAfter=30
    )
    story.append(Paragraph(f'Outage Report - {region}', title_style))
    story.append(Spacer(1, 0.5 * cm))
    
    # Date and generation info
    info_style = styles['Normal']
    story.append(Paragraph(f'<b>Report Period:</b> {start_date} to {end_date}', info_style))
    story.append(Paragraph(f'<b>Generated:</b> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', info_style))
    story.append(Spacer(1, 0.5 * cm))
    
    # Summary metrics table
    story.append(Paragraph('<b>Executive Summary</b>', styles['Heading2']))
    summary_data = [
        ['Metric', 'Value'],
        ['Number of Outages', str(summary_stats.get('num_outages', 0))],
        ['Total Outage Minutes', f"{float(summary_stats.get('total_minutes', 0)):.2f}"],
        ['Total Outage Hours', f"{float(summary_stats.get('total_minutes', 0)) / 60:.2f}"],
        ['Average Duration (min)', f"{float(summary_stats.get('avg_duration_min', 0)):.2f}"],
        ['Total Load Loss (MWh)', f"{float(summary_stats.get('total_load_loss_mwh', 0)):.2f}"]
    ]
    
    summary_table = Table(summary_data, colWidths=[8 * cm, 5 * cm])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 0.5 * cm))
    
    # Station summary table
    if not station_summary.empty:
        story.append(Paragraph('<b>Top Stations by Outage Minutes</b>', styles['Heading2']))
        station_data = [['Station', 'Outages', 'Total Min', 'Avg Min', 'Total Load Loss (MWh)', 'Avg Load Loss (MWh)']]
        for _, row in station_summary.head(10).iterrows():
            station_data.append([
                str(row.get('station', '')),
                str(row.get('outages_count', '')),
                f"{float(row.get('total_outage_min', 0)):.2f}",
                f"{float(row.get('avg_outage_min', 0)):.2f}",
                f"{float(row.get('total_load_loss_mwh', 0)):.2f}",
                f"{float(row.get('avg_load_loss_mwh', 0)):.2f}"
            ])
        
        station_table = Table(station_data, colWidths=[6 * cm, 3 * cm, 4 * cm, 4 * cm])
        station_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(station_table)
        story.append(Spacer(1, 0.5 * cm))
    
    # Feeder summary table
    if not feeder_summary.empty:
        story.append(Paragraph('<b>Top Feeders by Outage Hours</b>', styles['Heading2']))
        feeder_data = [['Feeder', 'Outages', 'Total Hours', 'Avg Hours', 'Total Load Loss (MWh)', 'Avg Load Loss (MWh)']]
        for _, row in feeder_summary.head(10).iterrows():
            feeder_data.append([
                str(row.get('feeder_33kv', '')),
                str(row.get('outages_count', '')),
                f"{float(row.get('outage_hrs', 0)):.2f}",
                f"{float(row.get('avg_outage_hrs', 0)):.2f}",
                f"{float(row.get('total_load_loss_mwh', 0)):.2f}",
                f"{float(row.get('avg_load_loss_mwh', 0)):.2f}"
            ])
        feeder_table = Table(feeder_data, colWidths=[5 * cm, 2.5 * cm, 3.5 * cm, 3.5 * cm, 4 * cm, 4 * cm])
        feeder_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(feeder_table)
        story.append(Spacer(1, 0.5 * cm))
    
    # Feeder-Party Pivot Table (Outage Table)
    if feeder_party_pivot is not None and not feeder_party_pivot.empty:
        story.append(Paragraph('<b>Outage Table - Feeder by Party Responsible</b>', styles['Heading2']))
        
        # Limit columns to avoid too many in PDF
        display_pivot = feeder_party_pivot.head(10)
        pivot_data = [list(display_pivot.columns)]
        for _, row in display_pivot.iterrows():
            formatted_row = []
            for val in row.values:
                if isinstance(val, (int, float)):
                    formatted_row.append(f"{float(val):.2f}")
                else:
                    formatted_row.append(str(val)[:15])
            pivot_data.append(formatted_row)
        
        pivot_table = Table(pivot_data, colWidths=[2.5 * cm] * min(len(display_pivot.columns), 6))
        pivot_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#70AD47')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(pivot_table)
        story.append(Spacer(1, 0.5 * cm))
    
    # SLA Violations Table
    if feeder_party_violations is not None and not feeder_party_violations.empty:
        story.append(Paragraph('<b>SLA Violations - TCN Negative Availability</b>', styles['Heading2']))
        
        display_violations = feeder_party_violations.head(10)
        violation_cols = ['station', 'feeder_33kv', 'TCN', 'max_hours_tcn', 'available_outage_hours_tcn']
        available_violation_cols = [c for c in violation_cols if c in display_violations.columns]
        
        violation_data = [available_violation_cols]
        for _, row in display_violations.iterrows():
            formatted_violation_row = []
            for col in available_violation_cols:
                val = row.get(col, '')
                if isinstance(val, (int, float)):
                    formatted_violation_row.append(f"{float(val):.2f}")
                else:
                    formatted_violation_row.append(f"{val}")
            violation_data.append(formatted_violation_row)
        
        violation_table = Table(violation_data, colWidths=[3 * cm] * len(available_violation_cols))
        violation_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#C5504A')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(violation_table)
        story.append(Spacer(1, 0.5 * cm))
    
    # Add chart images if provided
    if image_paths:
        story.append(PageBreak())
        story.append(Paragraph('<b>Charts & Visualizations</b>', styles['Heading2']))
        
        for img_path in image_paths:
            if os.path.exists(img_path):
                try:
                    story.append(RLImage(img_path, width=18 * cm, height=10 * cm))
                    story.append(Spacer(1, 0.5 * cm))
                except Exception:
                    pass
    
    # Build PDF
    doc.build(story)
    return output_path
