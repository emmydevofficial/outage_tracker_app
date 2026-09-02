"""Narrative Daily Outage Report generator (330kV/132kV TCN app).

Same "short story" template approach as
outage_tracker-main/utils/narrative_report.py, deliberately duplicated
rather than shared since the two apps are independent deployments with
their own column names (this app splits time into Hour_Off/Minute_Off
instead of a single time column, uses Substation/Equipment instead of
station/feeder, etc.) -- see that module's docstring for the design
rationale (dynamic/template-only, no LLM).
"""
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _fmt_hm(hour, minute) -> str:
    try:
        h, m = int(hour), int(minute)
        return f"{h:02d}:{m:02d}"
    except (TypeError, ValueError):
        return ""


def _clean(value) -> str:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    text = str(value).strip()
    return "" if text.lower() in ("nan", "none", "nat", "") else text


def _fmt_load(value) -> str:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    try:
        return f"{float(value):.1f}MW"
    except (TypeError, ValueError):
        return ""


def _is_still_open(row) -> bool:
    return not _clean(row.get("Date_On")) or pd.isna(row.get("Hour_On")) or pd.isna(row.get("Minute_On"))


def format_event_bullet(row) -> str:
    """One narrative sentence for a single outage event/row.

    Wording depends on Class: a Forced outage is an unplanned protection
    trip ("tripped due to X... Load lost: Y"), while an Emergency or
    Planned outage is a deliberate switching action ("was opened due to
    X... Load interrupted: Y") -- these are not interchangeable in TCN's
    own reporting convention.
    """
    still_open = _is_still_open(row)
    time_off = _fmt_hm(row.get("Hour_Off"), row.get("Minute_Off"))
    time_on = _fmt_hm(row.get("Hour_On"), row.get("Minute_On"))

    if time_off and time_on and not still_open:
        time_clause = f"{time_off}–{time_on} hrs"
    elif time_off:
        time_clause = f"{time_off} hrs"
    else:
        time_clause = ""

    substation = _clean(row.get("Substation"))
    equipment = _clean(row.get("Equipment"))
    line_label = " – ".join(x for x in [substation, equipment] if x) or "A line"

    is_forced = _clean(row.get("Class")).lower() == "forced"
    cause = _clean(row.get("Event_Indication"))
    if is_forced:
        cause_clause = f"tripped due to {cause}" if cause else "tripped"
    else:
        cause_clause = f"was opened due to {cause}" if cause else "was opened"

    lead = f"{time_clause}: {line_label} {cause_clause}." if time_clause else f"{line_label} {cause_clause}."
    clauses = [lead]

    load = _fmt_load(row.get("Last_Load_MW"))
    if load:
        load_label = "Load lost" if is_forced else "Load interrupted"
        clauses.append(f"{load_label}: {load}.")

    weather = _clean(row.get("Weather_Condition"))
    if weather:
        clauses.append(f"{weather} weather.")

    remarks = _clean(row.get("Remarks"))
    if remarks:
        clauses.append(remarks if remarks.endswith((".", "!", "?")) else f"{remarks}.")

    if still_open:
        clauses.append("Still out at time of report.")
    else:
        officer = _clean(row.get("Officer_Restoration"))
        if officer:
            clauses.append(f"Restored — confirmed by {officer}.")

    return " ".join(clauses)


def build_region_sections(df: pd.DataFrame) -> dict:
    """Group events into {region: {substation: [bullet, ...]}}, station/time ordered."""
    sections: dict = {}
    if df.empty:
        return sections
    df = df.sort_values(["Region", "Substation", "Hour_Off", "Minute_Off"])
    for region, region_df in df.groupby("Region", dropna=False):
        region_name = _clean(region) or "Unassigned Region"
        stations: dict = {}
        for substation, station_df in region_df.groupby("Substation", dropna=False):
            station_name = _clean(substation) or "Unnamed Substation"
            stations[station_name] = [format_event_bullet(r) for _, r in station_df.iterrows()]
        sections[region_name] = stations
    return sections


def score_notability(row) -> float:
    """Rank events for the highlights section -- bigger load lost, longer
    duration, and still-open events score higher."""
    score = 0.0
    try:
        load = float(row.get("Last_Load_MW"))
        if not pd.isna(load):
            score += load
    except (TypeError, ValueError):
        pass

    duration = row.get("Duration")
    try:
        duration_val = float(duration)
        if not pd.isna(duration_val):
            score += duration_val * 2
    except (TypeError, ValueError):
        if _clean(duration):
            score += 5

    if _is_still_open(row):
        score += 100

    if _clean(row.get("Class")).lower() in ("forced", "emergency"):
        score += 10

    return score


def build_highlights(df: pd.DataFrame, top_n: int = 4) -> list:
    if df.empty:
        return []
    scored = df.copy()
    scored["_score"] = scored.apply(score_notability, axis=1)
    scored = scored.sort_values("_score", ascending=False).head(top_n)
    return [format_event_bullet(r) for _, r in scored.iterrows()]


def _dominant_issue(df: pd.DataFrame) -> str:
    if df.empty:
        return "No significant system issues reported for this date."
    causes = df["Event_Indication"].dropna().astype(str).str.strip()
    causes = causes[causes != ""]
    if causes.empty:
        return "No significant system issues reported for this date."
    counts = causes.value_counts()
    top_cause, count = counts.index[0], int(counts.iloc[0])
    return f"{top_cause} was the most recurring issue, accounting for {count} of {len(df)} event(s) on this date."


def _grid_status_at_close(open_df: pd.DataFrame) -> str:
    if open_df.empty:
        return "All outages recorded had been restored by the close of this report."
    return f"{len(open_df)} outage(s) remained open at the close of this report, carried forward until restoration."


def build_at_a_glance(df: pd.DataFrame, open_df: pd.DataFrame) -> dict:
    class_counts = {}
    regions_covered = []
    if not df.empty:
        class_counts = df["Class"].fillna("Unclassified").value_counts().to_dict()
        regions_covered = sorted({_clean(r) for r in df["Region"] if _clean(r)})

    return {
        "total_events": len(df),
        "class_counts": class_counts,
        "regions_covered": regions_covered,
        "still_open_count": len(open_df),
        "lowest_frequency": "N/A — not currently recorded",
        "new_assets_energised": "N/A — not currently recorded",
        "dominant_issue": _dominant_issue(df),
        "grid_status_at_close": _grid_status_at_close(open_df),
    }


def build_still_open_section(open_df: pd.DataFrame) -> list:
    if open_df.empty:
        return []
    open_df = open_df.sort_values(["Date_Off", "Hour_Off", "Minute_Off"])
    lines = []
    for _, row in open_df.iterrows():
        substation = _clean(row.get("Substation"))
        equipment = _clean(row.get("Equipment"))
        label = " – ".join(x for x in [substation, equipment] if x) or "A line"
        date_off = _clean(row.get("Date_Off"))
        time_off = _fmt_hm(row.get("Hour_Off"), row.get("Minute_Off"))
        since = f"{date_off} {time_off}".strip()
        cause = _clean(row.get("Event_Indication"))
        cause_clause = f" due to {cause}" if cause else ""
        lines.append(f"{label} has been out since {since}{cause_clause}.")
    return lines


def generate_narrative_word_report(
    report_date: str,
    at_a_glance: dict,
    highlights: list,
    region_sections: dict,
    still_open_lines: list,
    output_path: str = None,
    subtitle: str = None,
) -> str:
    """Build the prose-style Daily Outage Report Word doc, mirroring
    outage_tracker-main/utils/report_generator.py::generate_narrative_word_report
    (same heading/paragraph conventions) since this app has no
    report_generator.py of its own."""
    if output_path is None:
        output_path = f"Daily_Outage_Report_{report_date}.docx"

    doc = Document()

    title = doc.add_heading(f'Daily Outage Report — {report_date}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        sub_para = doc.add_paragraph(subtitle)
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_date = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('At a Glance', 1)
    class_counts = at_a_glance.get('class_counts', {})
    class_line = ", ".join(f"{k}: {v}" for k, v in class_counts.items()) or "No events recorded"
    doc.add_paragraph(f"Total Events: {at_a_glance.get('total_events', 0)} ({class_line})")
    regions_covered = at_a_glance.get('regions_covered', [])
    doc.add_paragraph(f"Regions Covered: {', '.join(regions_covered) if regions_covered else 'None'}")
    doc.add_paragraph(f"Still-Open Outages (all dates): {at_a_glance.get('still_open_count', 0)}")
    doc.add_paragraph(f"Lowest Recorded Frequency: {at_a_glance.get('lowest_frequency', 'N/A')}")
    doc.add_paragraph(f"New Assets Energised: {at_a_glance.get('new_assets_energised', 'N/A')}")
    if at_a_glance.get('dominant_issue'):
        doc.add_paragraph(f"Dominant System Issue: {at_a_glance['dominant_issue']}")
    if at_a_glance.get('grid_status_at_close'):
        doc.add_paragraph(f"Grid Status at Close: {at_a_glance['grid_status_at_close']}")

    doc.add_heading('System Observations', 1)
    if highlights:
        for h in highlights:
            doc.add_paragraph(h, style='List Bullet')
    else:
        doc.add_paragraph("No notable events recorded for this date.")

    doc.add_heading('Region-by-Region Summary', 1)
    if region_sections:
        for region, stations in region_sections.items():
            doc.add_heading(region, 2)
            for station, bullets in stations.items():
                p = doc.add_paragraph()
                p.add_run(station).bold = True
                for bullet in bullets:
                    doc.add_paragraph(bullet, style='List Bullet')
    else:
        doc.add_paragraph("No outage events recorded for this date.")

    doc.add_heading('Outages Still Open (Carried Over)', 1)
    if still_open_lines:
        for line in still_open_lines:
            doc.add_paragraph(line, style='List Bullet')
    else:
        doc.add_paragraph("No outages currently open.")

    doc.save(output_path)
    return output_path


def generate_narrative_pdf_report(
    report_date: str,
    at_a_glance: dict,
    highlights: list,
    region_sections: dict,
    still_open_lines: list,
    output_path: str = None,
    subtitle: str = None,
) -> str:
    """PDF rendition of the Daily Outage Report, laid out like the source
    TCC report this feature was modeled on -- mirrors
    outage_tracker-main/utils/report_generator.py::generate_narrative_pdf_report
    (same layout/styling) since this app has no report_generator.py of its
    own."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib import colors
    from xml.sax.saxutils import escape as _esc

    if output_path is None:
        output_path = f"Daily_Outage_Report_{report_date}.pdf"

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm, leftMargin=1.8 * cm, rightMargin=1.8 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'NarrativeTitle', parent=styles['Heading1'], fontSize=17, alignment=TA_CENTER,
        textColor=colors.HexColor('#1a3a6b'), spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        'NarrativeSubtitle', parent=styles['Normal'], fontSize=9.5, alignment=TA_CENTER,
        textColor=colors.HexColor('#666666'), spaceAfter=4,
    )
    sub_heading_style = ParagraphStyle(
        'SubHeading', parent=styles['Heading2'], fontSize=13,
        textColor=colors.HexColor('#1a3a6b'), spaceBefore=8, spaceAfter=6,
    )
    station_style = ParagraphStyle(
        'StationHeading', parent=styles['Heading3'], fontSize=10.5,
        textColor=colors.HexColor('#1a3a6b'), spaceBefore=8, spaceAfter=2,
    )
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13, spaceAfter=4)
    bullet_style = ParagraphStyle('Bullet', parent=body_style, leftIndent=14)

    story = [
        Paragraph(f"SUMMARY OF OUTAGES &amp; EVENTS — {report_date}", title_style),
    ]
    if subtitle:
        story.append(Paragraph(_esc(subtitle), subtitle_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", subtitle_style))
    story.append(Spacer(1, 0.4 * cm))

    class_counts = at_a_glance.get('class_counts', {})
    regions_covered = at_a_glance.get('regions_covered', [])
    glance_rows = [
        ["Report Date", report_date],
        ["Issuing Authority", "Transmission Company of Nigeria (TCN)"],
        ["Regions Covered", ", ".join(regions_covered) if regions_covered else "None"],
        ["Forced Events", str(class_counts.get("Forced", 0))],
        ["Emergency Events", str(class_counts.get("Emergency", 0))],
        ["Planned Events", str(class_counts.get("Planned", 0))],
        ["Total Events", str(at_a_glance.get("total_events", 0))],
        ["Still-Open Outages (all dates)", str(at_a_glance.get("still_open_count", 0))],
        ["Lowest Recorded Frequency", at_a_glance.get("lowest_frequency", "N/A")],
        ["New Assets Energised", at_a_glance.get("new_assets_energised", "N/A")],
    ]
    glance_table = Table(glance_rows, colWidths=[6.3 * cm, 10.4 * cm])
    glance_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9.5),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ROWBACKGROUNDS', (1, 0), (1, -1), [colors.white, colors.HexColor('#F4F6FB')]),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    story.append(glance_table)
    story.append(Spacer(1, 0.4 * cm))

    if at_a_glance.get('dominant_issue'):
        story.append(Paragraph("<b>Dominant System Issue</b>", station_style))
        story.append(Paragraph(_esc(at_a_glance['dominant_issue']), body_style))
    if at_a_glance.get('grid_status_at_close'):
        story.append(Paragraph("<b>Grid Status at Close</b>", station_style))
        story.append(Paragraph(_esc(at_a_glance['grid_status_at_close']), body_style))

    story.append(PageBreak())
    story.append(Paragraph("REGION-BY-REGION SUMMARY", sub_heading_style))
    if region_sections:
        for region, stations in region_sections.items():
            region_bar = Table([[region.upper()]], colWidths=[16.9 * cm])
            region_bar.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#1a3a6b')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.whitesmoke),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(Spacer(1, 0.3 * cm))
            story.append(region_bar)
            for station, bullets in stations.items():
                story.append(Paragraph(_esc(station), station_style))
                for bullet in bullets:
                    story.append(Paragraph(f"• {_esc(bullet)}", bullet_style))
    else:
        story.append(Paragraph("No outage events recorded for this date.", body_style))

    story.append(PageBreak())
    story.append(Paragraph("SYSTEM OBSERVATIONS", sub_heading_style))
    if highlights:
        for i, h in enumerate(highlights, start=1):
            story.append(Paragraph(f"<b>Notable Event {i}:</b> {_esc(h)}", bullet_style))
    else:
        story.append(Paragraph("No notable events recorded for this date.", body_style))

    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("OUTAGES STILL OPEN (CARRIED OVER)", sub_heading_style))
    if still_open_lines:
        for line in still_open_lines:
            story.append(Paragraph(f"• {_esc(line)}", bullet_style))
    else:
        story.append(Paragraph("No outages currently open.", body_style))

    doc.build(story)
    return output_path
